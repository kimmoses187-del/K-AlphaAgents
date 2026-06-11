"""
report/exporters.py
===================
Gap 7 — Excel and Word export of analysis results.

Produces two optional output artefacts alongside the existing PDF:
  1. {date}_portfolio.xlsx  — portfolio weights, conviction scores, backtest
                               summary on separate sheets
  2. {date}_report.docx     — Word document converted from the per-stock
                               markdown reports (readable by non-technical users)

Both are saved to the same output directory as the PDF.

Public API
----------
from report.exporters import export_portfolio_xlsx, export_reports_docx

xlsx_path = export_portfolio_xlsx(portfolios, backtest_results,
                                   company_names, as_of_date, output_dir)
docx_path = export_reports_docx(report_md_paths, company_names,
                                  as_of_date, output_dir)
"""

from __future__ import annotations

import logging
import os
import re
from datetime import datetime
from typing import Dict, List, Optional

log = logging.getLogger(__name__)


# ── Excel Export ──────────────────────────────────────────────────────────────

def export_portfolio_xlsx(
    portfolios: Dict,
    backtest_results: Optional[Dict],
    company_names: Dict[str, str],
    as_of_date: datetime,
    output_dir: str,
) -> Optional[str]:
    """
    Export portfolio allocation and backtest summary to Excel.

    Sheets:
      1. Portfolio Summary  — signal, conviction, weight per stock per profile
      2. Backtest Summary   — total return, Sharpe, max drawdown per profile
      3. Signal Details     — debate consensus type, rounds per stock per profile

    Returns path to saved .xlsx, or None on failure.
    """
    try:
        import openpyxl
        from openpyxl.styles import (Font, PatternFill, Alignment,
                                      Border, Side, GradientFill)
        from openpyxl.utils import get_column_letter
    except ImportError:
        log.warning("openpyxl not installed — skipping Excel export")
        return None

    # ── Colour palette ────────────────────────────────────────────────────
    NAVY  = "0D1117"
    GOLD  = "F0B429"
    GREEN = "2EA043"
    RED   = "F85149"
    WHITE = "E6EDF3"
    GREY  = "8B949E"
    PANEL = "161B22"

    def _header_fill(hex_color: str) -> PatternFill:
        return PatternFill("solid", fgColor=hex_color)

    def _cell_font(bold=False, color=WHITE) -> Font:
        return Font(name="Calibri", size=10, bold=bold, color=color)

    def _thin_border() -> Border:
        s = Side(border_style="thin", color="2D333B")
        return Border(left=s, right=s, top=s, bottom=s)

    def _style_header_row(ws, row_num: int, fill_color: str, n_cols: int):
        for col in range(1, n_cols + 1):
            cell = ws.cell(row=row_num, column=col)
            cell.fill   = _header_fill(fill_color)
            cell.font   = _cell_font(bold=True, color=WHITE)
            cell.border = _thin_border()
            cell.alignment = Alignment(horizontal="center", vertical="center")

    def _style_data_row(ws, row_num: int, n_cols: int, alt: bool = False):
        fill = _header_fill(PANEL if alt else NAVY)
        for col in range(1, n_cols + 1):
            cell = ws.cell(row=row_num, column=col)
            cell.fill   = fill
            cell.font   = _cell_font()
            cell.border = _thin_border()
            cell.alignment = Alignment(horizontal="center", vertical="center")

    wb = openpyxl.Workbook()

    # Profiles actually present in this run (single- or multi-profile)
    PROFILES = list(portfolios.keys())

    # ── Sheet 1: Portfolio Summary ─────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Portfolio Summary"
    ws1.sheet_view.showGridLines = False

    headers = ["Profile", "Stock Code", "Company", "Signal", "Convergence", "Weight (%)"]
    for col, h in enumerate(headers, start=1):
        ws1.cell(row=1, column=col, value=h)
    _style_header_row(ws1, 1, GOLD, len(headers))
    # Gold header text should be dark
    for col in range(1, len(headers) + 1):
        ws1.cell(row=1, column=col).font = Font(name="Calibri", size=10, bold=True, color=NAVY)

    row = 2
    for profile in PROFILES:
        po = portfolios.get(profile, {})
        allocs = po.get("stock_allocations", {})
        for code, alloc in allocs.items():
            signal     = alloc.get("signal", "N/A")
            conviction = alloc.get("conviction", 0)
            weight     = alloc.get("weight", 0) * 100
            ws1.cell(row=row, column=1, value=profile.title())
            ws1.cell(row=row, column=2, value=code)
            ws1.cell(row=row, column=3, value=company_names.get(code, code))
            ws1.cell(row=row, column=4, value=signal)
            ws1.cell(row=row, column=5, value=round(conviction, 3))
            ws1.cell(row=row, column=6, value=round(weight, 1))
            _style_data_row(ws1, row, len(headers), alt=(row % 2 == 0))
            # Colour-code signal cell
            sig_cell = ws1.cell(row=row, column=4)
            sig_cell.fill = _header_fill(GREEN if signal == "BUY" else RED)
            sig_cell.font = _cell_font(bold=True)
            row += 1

    # Column widths
    for col, width in enumerate([16, 12, 30, 8, 12, 12], start=1):
        ws1.column_dimensions[get_column_letter(col)].width = width

    # ── Sheet 2: Backtest Summary ──────────────────────────────────────────
    ws2 = wb.create_sheet("Backtest Summary")
    ws2.sheet_view.showGridLines = False

    b_headers = ["Profile", "Total Return (%)", "Sharpe Ratio", "Max Drawdown (%)", "vs Benchmark (%)"]
    for col, h in enumerate(b_headers, start=1):
        ws2.cell(row=1, column=col, value=h)
    _style_header_row(ws2, 1, GOLD, len(b_headers))
    for col in range(1, len(b_headers) + 1):
        ws2.cell(row=1, column=col).font = Font(name="Calibri", size=10, bold=True, color=NAVY)

    br = 2
    if backtest_results:
        for profile in PROFILES:
            engine = backtest_results.get(profile)
            if engine is None:
                continue
            summary = getattr(engine, "summary", {}) or {}
            # Try to extract from engine results
            try:
                agent_results = engine.results
                if agent_results and "Agent Portfolio" in agent_results:
                    ap = agent_results["Agent Portfolio"]
                    total_ret = round(((ap.get("cumulative_return") or pd.Series([1])).iloc[-1] - 1) * 100, 2)
                else:
                    total_ret = "N/A"
            except Exception:
                total_ret = "N/A"

            ws2.cell(row=br, column=1, value=profile.title())
            ws2.cell(row=br, column=2, value=total_ret)
            ws2.cell(row=br, column=3, value="see PDF")
            ws2.cell(row=br, column=4, value="see PDF")
            ws2.cell(row=br, column=5, value="see PDF")
            _style_data_row(ws2, br, len(b_headers), alt=(br % 2 == 0))
            br += 1
    else:
        ws2.cell(row=2, column=1, value="No backtest run (all SELL or not requested)")
        _style_data_row(ws2, 2, len(b_headers))

    for col, width in enumerate([16, 18, 14, 18, 18], start=1):
        ws2.column_dimensions[get_column_letter(col)].width = width

    # ── Sheet 3: Debate Details ────────────────────────────────────────────
    # (imported from all_results if passed — lightweight metadata only)
    ws3 = wb.create_sheet("Signal Details")
    ws3.sheet_view.showGridLines = False

    d_headers = ["Profile", "Stock Code", "Company", "Final Signal",
                 "Consensus Type", "Debate Rounds", "Analysis Date"]
    for col, h in enumerate(d_headers, start=1):
        ws3.cell(row=1, column=col, value=h)
    _style_header_row(ws3, 1, GOLD, len(d_headers))
    for col in range(1, len(d_headers) + 1):
        ws3.cell(row=1, column=col).font = Font(name="Calibri", size=10, bold=True, color=NAVY)

    dr = 2
    date_tag = as_of_date.strftime("%Y-%m-%d")
    for profile in PROFILES:
        po = portfolios.get(profile, {})
        allocs = po.get("stock_allocations", {})
        for code, alloc in allocs.items():
            ws3.cell(row=dr, column=1, value=profile.title())
            ws3.cell(row=dr, column=2, value=code)
            ws3.cell(row=dr, column=3, value=company_names.get(code, code))
            ws3.cell(row=dr, column=4, value=alloc.get("signal", "N/A"))
            ws3.cell(row=dr, column=5, value="—")   # not stored in portfolio dict
            ws3.cell(row=dr, column=6, value="—")
            ws3.cell(row=dr, column=7, value=date_tag)
            _style_data_row(ws3, dr, len(d_headers), alt=(dr % 2 == 0))
            sig_cell = ws3.cell(row=dr, column=4)
            sig_cell.fill = _header_fill(GREEN if alloc.get("signal") == "BUY" else RED)
            sig_cell.font = _cell_font(bold=True)
            dr += 1

    for col, width in enumerate([16, 12, 30, 14, 16, 14, 14], start=1):
        ws3.column_dimensions[get_column_letter(col)].width = width

    # ── Save ──────────────────────────────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"portfolio_{as_of_date.strftime('%Y-%m-%d')}.xlsx")
    wb.save(out_path)
    log.info("Portfolio Excel exported: %s", out_path)
    return out_path


# ── Word Export ───────────────────────────────────────────────────────────────

def export_reports_docx(
    report_md_paths: Dict[str, Dict[str, str]],
    company_names: Dict[str, str],
    as_of_date: datetime,
    output_dir: str,
) -> Optional[str]:
    """
    Convert per-stock markdown reports to a single Word document.

    Parameters
    ----------
    report_md_paths : {stock_code: {profile: md_file_path}}
    company_names   : {stock_code: company_name}
    as_of_date      : analysis date
    output_dir      : directory for output

    Returns path to saved .docx, or None on failure.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("python-docx not installed — skipping Word export")
        return None

    doc = Document()

    # ── Document styling ──────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title page
    title = doc.add_heading("K-AlphaAgents — Equity Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Analysis Date: {as_of_date.strftime('%Y-%m-%d')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    stocks_list = ", ".join(
        f"{name} ({code})" for code, name in company_names.items()
    )
    doc.add_paragraph(f"Stocks analysed: {stocks_list}")
    doc.add_page_break()

    # ── Per-stock sections ─────────────────────────────────────────────────
    def _md_to_docx(doc, md_text: str):
        """Naive markdown → docx converter for the subset used in reports."""
        for line in md_text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("#### "):
                doc.add_heading(stripped[5:], level=4)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("---"):
                doc.add_paragraph("─" * 50)
            elif stripped.startswith("|"):
                # Table row — just render as plain text for simplicity
                cells = [c.strip() for c in stripped.split("|") if c.strip()]
                doc.add_paragraph("  |  ".join(cells))
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped == "":
                doc.add_paragraph()
            else:
                # Strip bold/italic markers
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
                clean = re.sub(r"\*(.+?)\*", r"\1", clean)
                doc.add_paragraph(clean)

    for stock_code, profiles in report_md_paths.items():
        cname = company_names.get(stock_code, stock_code)
        doc.add_heading(f"{cname} ({stock_code})", level=1)

        for profile, md_path in profiles.items():
            doc.add_heading(f"Profile: {profile.title()}", level=2)
            if md_path and os.path.exists(md_path):
                try:
                    with open(md_path, encoding="utf-8") as f:
                        md_text = f.read()
                    _md_to_docx(doc, md_text)
                except Exception as exc:
                    doc.add_paragraph(f"[Error reading report: {exc}]")
            else:
                doc.add_paragraph("[Report file not found]")

            doc.add_page_break()

    # ── Save ──────────────────────────────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"report_{as_of_date.strftime('%Y-%m-%d')}.docx")
    doc.save(out_path)
    log.info("Word report exported: %s", out_path)
    return out_path
